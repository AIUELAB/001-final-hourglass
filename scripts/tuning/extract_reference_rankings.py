#!/usr/bin/env python3
"""
参照ランキング（Claude/Gemini docx）を構造化データに変換するスクリプト

出力形式:
    - JSON: 人物名、年齢、エピソードテキスト、ソース、順位
"""

import json
import re
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Optional

from docx import Document


@dataclass
class ReferenceEpisode:
    """参照エピソードデータ"""

    rank: int
    person_name: str
    age: Optional[int]
    episode_text: str
    source: str  # "claude" or "gemini"
    raw_line: str  # デバッグ用


def extract_claude_ranking(docx_path: str) -> list[ReferenceEpisode]:
    """
    Claude docxからランキングを抽出

    形式:
        第N位　人物名（年齢）
        エピソードテキスト
    """
    doc = Document(docx_path)
    episodes = []
    current_rank = None
    current_name = None
    current_age = None

    # 「第N位　人物名（年齢）」のパターン
    rank_pattern = re.compile(r"第(\d+)位[　\s]+(.+?)(?:（|[\(])(\d+)歳(?:）|[\)])")

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # ランク行をチェック
        match = rank_pattern.match(text)
        if match:
            current_rank = int(match.group(1))
            current_name = match.group(2).strip()
            current_age = int(match.group(3))
        elif current_rank and current_name:
            # エピソードテキスト行
            if len(text) > 50:  # 短すぎるものはスキップ
                episodes.append(
                    ReferenceEpisode(
                        rank=current_rank,
                        person_name=current_name,
                        age=current_age,
                        episode_text=text,
                        source="claude",
                        raw_line=f"第{current_rank}位　{current_name}（{current_age}歳）",
                    )
                )
                current_rank = None
                current_name = None
                current_age = None

    return episodes


def extract_gemini_ranking(docx_path: str) -> list[ReferenceEpisode]:
    """
    Gemini docxからランキングを抽出

    形式:
        【第N位】人物名（職業など）
        年齢：N歳〜M歳 「タイトル」 エピソードテキスト
    """
    doc = Document(docx_path)
    episodes = []
    current_rank = None
    current_name = None

    # 「【第N位】人物名（職業）」のパターン - 人物名は全角括弧の前まで
    rank_pattern = re.compile(r"【第(\d+)位】([^（\(]+)")
    # 「年齢：N歳」のパターン
    age_pattern = re.compile(r"年齢[：:](\d+)歳(?:〜(\d+)歳)?")

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # ランク行をチェック
        match = rank_pattern.match(text)
        if match:
            current_rank = int(match.group(1))
            current_name = match.group(2).strip()
        elif current_rank and current_name:
            # 年齢とエピソード行
            age_match = age_pattern.search(text)
            if age_match:
                age_start = int(age_match.group(1))
                # エピソードテキストは年齢の後
                ep_text = text[age_match.end() :].strip()
                # 「タイトル」を除去
                ep_text = re.sub(r"^「[^」]+」\s*", "", ep_text)

                if len(ep_text) > 50:
                    episodes.append(
                        ReferenceEpisode(
                            rank=current_rank,
                            person_name=current_name,
                            age=age_start,
                            episode_text=ep_text,
                            source="gemini",
                            raw_line=f"【第{current_rank}位】{current_name}",
                        )
                    )
                current_rank = None
                current_name = None

    return episodes


def main():
    """メイン処理"""
    output_dir = Path("/Users/admin/Documents/AIUELAB/001-final-hourglass/src/reports")
    output_dir.mkdir(parents=True, exist_ok=True)

    # Claude docx抽出
    claude_path = "/Users/admin/Desktop/claude_有名人エピソードBest100.docx"
    print(f"Claude docxを読み込み中: {claude_path}")
    claude_eps = extract_claude_ranking(claude_path)
    print(f"  抽出件数: {len(claude_eps)}")

    # Gemini docx抽出
    gemini_path = "/Users/admin/Desktop/Gemini.docx"
    print(f"Gemini docxを読み込み中: {gemini_path}")
    gemini_eps = extract_gemini_ranking(gemini_path)
    print(f"  抽出件数: {len(gemini_eps)}")

    # 結果保存
    result = {
        "claude": [asdict(e) for e in claude_eps],
        "gemini": [asdict(e) for e in gemini_eps],
        "summary": {
            "claude_count": len(claude_eps),
            "gemini_count": len(gemini_eps),
            "claude_persons": list({e.person_name for e in claude_eps}),
            "gemini_persons": list({e.person_name for e in gemini_eps}),
        },
    }

    output_path = output_dir / "reference_rankings_extracted.json"
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(result, f, ensure_ascii=False, indent=2)

    print(f"\n結果を保存: {output_path}")

    # サンプル表示
    print("\n=== Claude Top 10 ===")
    for e in claude_eps[:10]:
        print(f"  {e.rank}. {e.person_name}（{e.age}歳）: {e.episode_text[:50]}...")

    print("\n=== Gemini Top 10 ===")
    for e in gemini_eps[:10]:
        print(f"  {e.rank}. {e.person_name}（{e.age}歳）: {e.episode_text[:50]}...")


if __name__ == "__main__":
    main()
